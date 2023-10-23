import tkinter as tk
from docx import Document
import os
import shutil
from datetime import date

def buyukharf(cumle):
    kelimeler = cumle.split()
    duzeltilmis_kelimeler = [kelime.capitalize() for kelime in kelimeler]
    duzeltilmis_cumle = " ".join(duzeltilmis_kelimeler)
    return duzeltilmis_cumle

def ihalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı"]:
            ek = 'ı'
        elif onceki_harf in ["e", "i"]:
            ek = 'i'
        elif onceki_harf in ["o", "u"]:
            ek = 'u'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ü'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı"]:
            ek = 'ı'
        elif onceki_harf in ["e", "i"]:
            ek = 'i'
        elif onceki_harf in ["o", "u"]:
            ek = 'u'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ü'
    elif harfler == 's':
        if son_harf in ["a", "ı"]:
            ek = 'yı'
        elif son_harf in ["e", "i"]:
            ek = 'yi'
        elif son_harf in ["o", "u"]:
            ek = 'yu'
        elif son_harf in ["ö", "ü"]:
            ek = 'yü'

    return f"{buyukharf(kelime)}'{ek}"

def ehalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'a'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'e'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'a'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'e'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'ya'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'ye'

    return f"{buyukharf(kelime)}'{ek}"


def dehalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'

    return f"{buyukharf(kelime)}'{ek}"

def denhalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'dan'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'den'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'dan'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'den'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'dan'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'den'

    return f"{buyukharf(kelime)}'{ek}"

def ninhalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

   

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı"]:
            ek = 'ın'
        elif onceki_harf in ["e", "i"]:
            ek = 'in'
        elif onceki_harf in ["o", "u"]:
            ek = 'un'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ün'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı"]:
            ek = 'ın'
        elif onceki_harf in ["e", "i"]:
            ek = 'in'
        elif onceki_harf in ["o", "u"]:
            ek = 'un'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ün'
    elif harfler == 's':
        if son_harf in ["a", "ı"]:
            ek = 'nın'
        elif son_harf in ["e", "i"]:
            ek = 'nin'
        elif son_harf in ["o", "u"]:
            ek = 'nun'
        elif son_harf in ["ö", "ü"]:
            ek = 'nün'

    return f"{buyukharf(kelime)}'{ek}"

def debhalii(kelime): #de-da bağlacı
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'

    return f"{buyukharf(kelime)} {ek}"



# Bu işlev, verileri txt dosyasından alır ve form alanlarına yansıtır
def load_data_from_txt(ad_entry, hikayeler_entry, cinsiyet_entry, tip_entry, ihali_entry, ehali_entry, dehali_entry, denhali_entry, debhali_entry, ninhali_entry):
    with open("siparis.txt", "r") as file:
        lines = file.readlines()

    veriler = {}
    for line in lines:
        key, value = line.strip().split(":")
        veriler[key] = value

    ad_entry.delete(0, tk.END)
    ad_entry.insert(0, veriler.get("ad", ""))
    hikayeler_entry.delete(0, tk.END)
    hikayeler_entry.insert(0, veriler.get("hikayeler", ""))
    cinsiyet_entry.delete(0, tk.END)
    cinsiyet_entry.insert(0, veriler.get("cinsiyet", ""))
    tip_entry.delete(0, tk.END)
    tip_entry.insert(0, veriler.get("tip", ""))
    ad = veriler.get("ad", "")
    ihali = ihalii(ad)
    ehali = ehalii(ad)
    dehali = dehalii(ad)
    denhali = denhalii(ad)
    debhali = debhalii(ad)
    ninhali = ninhalii(ad)
    ihali_entry.delete(0, tk.END)
    ihali_entry.insert(0, ihali)
    ehali_entry.delete(0, tk.END)
    ehali_entry.insert(0, ehali)
    dehali_entry.delete(0, tk.END)
    dehali_entry.insert(0, dehali)
    denhali_entry.delete(0, tk.END)
    denhali_entry.insert(0, denhali)
    debhali_entry.delete(0, tk.END)
    debhali_entry.insert(0, debhali)
    ninhali_entry.delete(0, tk.END)
    ninhali_entry.insert(0, ninhali)


   
# Bu işlev, verileri form alanlarından alır ve belgeyi üretir
def generate_docs(ad_entry, hikayeler_entry, cinsiyet_entry, tip_entry, ihali_entry, ehali_entry, dehali_entry, denhali_entry, debhali_entry, ninhali_entry):
    ad = ad_entry.get()
    hikayeler = hikayeler_entry.get()
    cinsiyet = cinsiyet_entry.get()
    tip = tip_entry.get()
    tarih = date.today().strftime("%Y-%m-%d")
    hikaye_numaralari = hikayeler.split(',')
    hikaye_isimleri = ["hikaye-" + num for num in hikaye_numaralari]

    hikayalar_dizini = os.path.join("hikayeler", cinsiyet, tip)

   # ... (previous code) ...

    for hikaye_adi in hikaye_isimleri:
        original_path = os.path.join(hikayalar_dizini, hikaye_adi + ".docx")
        copy_path = os.path.join(hikayalar_dizini, hikaye_adi + "_" + ad + ".docx")
        shutil.copyfile(original_path, copy_path)
        
        ihali = ihalii(ad)
        ehali = ehalii(ad)
        dehali = dehalii(ad)
        denhali = denhalii(ad)
        debhali = debhalii(ad)
        ninhali = ninhalii(ad)
        word_replacements = {
            'xxdenxx': denhali,
            'xxi xx': ihali,
            'xxexx': ehali,
            'xxdexx': dehali,
            'yavuzdeb': debhali,
            'yavuzun': ninhali,
            'xxadxx': ad 
        }

        doc = Document(os.path.join(hikayalar_dizini, hikaye_adi + "_" + ad + ".docx"))
        new_doc = Document()

        for paragraph in doc.paragraphs:
            new_paragraph = new_doc.add_paragraph()
            for run in paragraph.runs:
                text = run.text
                for search_word, replace_word in word_replacements.items():
                    text = text.replace(search_word, replace_word)
                new_run = new_paragraph.add_run(text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

        new_doc.save(os.path.join("üretim", ad + "_" + hikaye_adi + "_" + tarih + ".docx"))
        result_label.config(text="Belgeler başarıyla üretildi!")
        # Orijinal dosyanın kopyasını sil
        os.remove(os.path.join(hikayalar_dizini, hikaye_adi + "_" + ad + ".docx"))

        # Sipariş dosyasını taşı ve adını değiştir
    #txt dosyasını taşı
    # yeni_ad = ad+"-" + tarih + ".txt"
    # shutil.move("siparis.txt", os.path.join("üretilmişler", yeni_ad))

# Tkinter penceresini oluştur
window = tk.Tk()
window.title("Belge Üretici")

# Form için bir Frame oluşturun
form_frame = tk.Frame(window,padx=100,pady=100)
form_frame.pack()

# Verileri Yükle butonunu oluştur
load_button = tk.Button(form_frame, text="Verileri Yükle", command=lambda: load_data_from_txt(
    labels_entries[0][1], labels_entries[1][1], labels_entries[2][1], labels_entries[3][1],
    labels_entries[4][1], labels_entries[5][1], labels_entries[6][1], labels_entries[7][1],
    labels_entries[8][1], labels_entries[9][1]))
load_button.grid(row=0, column=0, columnspan=2)

# Etiketler ve girdi alanları oluştur
labels_entries = [
    ("Ad:", tk.Entry(form_frame, width=30)),
    ("Hikayeler (Virgülle Ayırılmış):", tk.Entry(form_frame, width=30)),
    ("Cinsiyet:", tk.Entry(form_frame, width=30)),
    ("Tip:", tk.Entry(form_frame, width=30)),
    ("İhali:", tk.Entry(form_frame, width=30)),
    ("Ehali:", tk.Entry(form_frame, width=30)),
    ("Dehali:", tk.Entry(form_frame, width=30)),
    ("Denhali:", tk.Entry(form_frame, width=30)),
    ("Debhali:", tk.Entry(form_frame, width=30)),
    ("Ninhali:", tk.Entry(form_frame, width=30))
]

# Etiketler ve girdi alanları için döngü oluşturarak düzenleme
row_counter = 1
for label, entry in labels_entries:
    tk.Label(form_frame, text=label,pady=5).grid(row=row_counter, column=0)
    entry.grid(row=row_counter, column=1)
    row_counter += 1

# Üret butonunu oluştur
generate_button = tk.Button(form_frame, text="Üret", command=lambda: generate_docs(
    labels_entries[0][1], labels_entries[1][1], labels_entries[2][1], labels_entries[3][1],
    labels_entries[4][1], labels_entries[5][1], labels_entries[6][1], labels_entries[7][1],
    labels_entries[8][1], labels_entries[9][1]))
generate_button.grid(row=row_counter, column=0, columnspan=2)

# Sonuç etiketi
result_label = tk.Label(window, text="")
result_label.pack()

# Formun genişliğini artırın
form_frame.update_idletasks()
form_frame.config(width=form_frame.winfo_reqwidth() + 100)

window.mainloop()