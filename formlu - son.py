import tkinter as tk
from docx import Document
import os
import shutil
from datetime import date

def buyukharh(cumle):
    kelimeler = cumle.split()
    duzeltilmis_kelimeler = [kelime.capitalize() for kelime in kelimeler]
    duzeltilmis_cumle = " ".join(duzeltilmis_kelimeler)
    return duzeltilmis_cumle

def ihalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı"]:
            ek = 'ı'
        elif onceki_harf in ["e", "i"]:
            ek = 'i'
        elif onceki_harf in ["o", "u"]:
            ek = 'u'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ü'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı"]:
            ek = 'ı'
        elif onceki_harf in ["e", "i"]:
            ek = 'i'
        elif onceki_harf in ["o", "u"]:
            ek = 'u'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ü'
    elif harfler == 's':
        if son_harf in ["a", "ı"]:
            ek = 'yı'
        elif son_harf in ["e", "i"]:
            ek = 'yi'
        elif son_harf in ["o", "u"]:
            ek = 'yu'
        elif son_harf in ["ö", "ü"]:
            ek = 'yü'

    return f"{kelime}'{ek}"

def ehalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'a'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'e'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'a'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'e'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'ya'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'ye'

    return f"{kelime}'{ek}"


def dehalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'

    return f"{kelime}'{ek}"

def denhalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'dan'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'den'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'dan'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'den'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'dan'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'den'

    return f"{kelime}'{ek}"

def ninhalii(kelime):
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

   

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı"]:
            ek = 'ın'
        elif onceki_harf in ["e", "i"]:
            ek = 'in'
        elif onceki_harf in ["o", "u"]:
            ek = 'un'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ün'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı"]:
            ek = 'ın'
        elif onceki_harf in ["e", "i"]:
            ek = 'in'
        elif onceki_harf in ["o", "u"]:
            ek = 'un'
        elif onceki_harf in ["ö", "ü"]:
            ek = 'ün'
    elif harfler == 's':
        if son_harf in ["a", "ı"]:
            ek = 'nın'
        elif son_harf in ["e", "i"]:
            ek = 'nin'
        elif son_harf in ["o", "u"]:
            ek = 'nun'
        elif son_harf in ["ö", "ü"]:
            ek = 'nün'

    return f"{kelime}'{ek}"

def debhalii(kelime): #de-da bağlacı
    kelime = kelime.lower()
    harfler = ""
    ek = ""
    son_harf = kelime[-1]
    son_iki_harf = kelime[-2] if len(kelime) >= 2 else ""

    print(f"{son_harf}\n")

    if son_harf in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "s"
    elif son_iki_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"] and son_harf not in ["a", "ı", "e", "i", "o", "u", "ö", "ü"]:
        harfler = "zz"
    else:
        harfler = "z"

    if harfler == 'zz':
        onceki_harf = kelime[-3] if len(kelime) >= 3 else ""
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 'z':
        onceki_harf = kelime[-2]
        if onceki_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif onceki_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'
    elif harfler == 's':
        if son_harf in ["a", "ı", "o", "u"]:
            ek = 'da'
        elif son_harf in ["e", "i", "ö", "ü"]:
            ek = 'de'

    return f"{kelime} {ek}"



# Bu işlev, verileri txt dosyasından alır ve form alanlarına yansıtır
def load_data_from_txt():
    with open("siparis.txt", "r") as file:
        lines = file.readlines()

    veriler = {}
    for line in lines:
        key, value = line.strip().split(":")
        veriler[key] = value

    entry_ad.delete(0, tk.END)
    entry_ad.insert(0, veriler.get("ad", ""))
    entry_hikayeler.delete(0, tk.END)
    entry_hikayeler.insert(0, veriler.get("hikayeler", ""))
    entry_cinsiyet.delete(0, tk.END)
    entry_cinsiyet.insert(0, veriler.get("cinsiyet", ""))
    entry_tip.delete(0, tk.END)
    entry_tip.insert(0, veriler.get("tip", ""))
    ad=veriler.get("ad", "")
    ihali = buyukharh(ihalii(ad))
    ehali = buyukharh(ehalii(ad))
    dehali = buyukharh(dehalii(ad))
    denhali = buyukharh(denhalii(ad))
    debhali = buyukharh(debhalii(ad))
    ninhali = buyukharh(ninhalii(ad))
    entry_ihali.delete(0, tk.END)
    entry_ihali.insert(0, ihali)
    entry_ehali.delete(0, tk.END)
    entry_ehali.insert(0, ehali)
    entry_dehali.delete(0, tk.END)
    entry_dehali.insert(0, dehali)
    entry_denhali.delete(0, tk.END)
    entry_denhali.insert(0, denhali)
    entry_debhali.delete(0, tk.END)
    entry_debhali.insert(0, debhali)
    entry_ninhali.delete(0, tk.END)
    entry_ninhali.insert(0, ninhali)


   
# Bu işlev, verileri form alanlarından alır ve belgeyi üretir
def generate_docs():
    ad = entry_ad.get()
    hikayeler = entry_hikayeler.get()
    cinsiyet = entry_cinsiyet.get()
    tip = entry_tip.get()
    tarih = date.today().strftime("%Y-%m-%d")
    hikaye_numaralari = hikayeler.split(',')
    hikaye_isimleri = ["hikaye-" + num for num in hikaye_numaralari]

    hikayalar_dizini = os.path.join("hikayeler", cinsiyet, tip)

   # ... (previous code) ...

    for hikaye_adi in hikaye_isimleri:
        original_path = os.path.join(hikayalar_dizini, hikaye_adi + ".docx")
        copy_path = os.path.join(hikayalar_dizini, hikaye_adi + "_" + ad + ".docx")
        shutil.copyfile(original_path, copy_path)

        ihali = buyukharh(ihalii(ad))
        ehali = buyukharh(ehalii(ad))
        dehali = buyukharh(dehalii(ad))
        denhali = buyukharh(denhalii(ad))
        debhali = buyukharh(debhalii(ad))
        ninhali = buyukharh(ninhalii(ad))
        word_replacements = {
            'xxdenxx': denhali,
            'xxi xx': ihali,
            'xxexx': ehali,
            'yavuzdan': dehali,
            'yavuzdeb': debhali,
            'yavuzun': ninhali,
            'xxadxx': ad 
        }

        doc = Document(os.path.join(hikayalar_dizini, hikaye_adi + "_" + ad + ".docx"))
        new_doc = Document()

        for paragraph in doc.paragraphs:
            new_paragraph = new_doc.add_paragraph()
            for run in paragraph.runs:
                text = run.text
                for search_word, replace_word in word_replacements.items():
                    text = text.replace(search_word, replace_word)
                new_run = new_paragraph.add_run(text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

        new_doc.save(os.path.join("üretim", ad + "_" + hikaye_adi + "_" + tarih + ".docx"))
        result_label.config(text="Belgeler başarıyla üretildi!")
        # Orijinal dosyanın kopyasını sil
        os.remove(os.path.join(hikayalar_dizini, hikaye_adi + "_" + ad + ".docx"))

        # Sipariş dosyasını taşı ve adını değiştir
    #txt dosyasını taşı
    # yeni_ad = ad+"-" + tarih + ".txt"
    # shutil.move("siparis.txt", os.path.join("üretilmişler", yeni_ad))

# Tkinter penceresini oluştur
window = tk.Tk()
window.title("Belge Üretici")

# Verileri Yükle butonunu oluştur
load_button = tk.Button(window, text="Verileri Yükle", command=load_data_from_txt)

# Etiketler ve girdi alanları oluştur
label_ad = tk.Label(window, text="Ad:")
entry_ad = tk.Entry(window)
label_hikayeler = tk.Label(window, text="Hikayeler (Virgülle Ayırılmış):")
entry_hikayeler = tk.Entry(window)
label_cinsiyet = tk.Label(window, text="Cinsiyet:")
entry_cinsiyet = tk.Entry(window)
label_tip = tk.Label(window, text="Tip:")
entry_tip = tk.Entry(window)
# Etiketler ve girdi alanları için ihali, ehali, dehali, denhali, debhali ve ninhali ekleyin
label_ihali = tk.Label(window, text="İhali:")
entry_ihali = tk.Entry(window)
label_ehali = tk.Label(window, text="Ehali:")
entry_ehali = tk.Entry(window)
label_dehali = tk.Label(window, text="Dehali:")
entry_dehali = tk.Entry(window)
label_denhali = tk.Label(window, text="Denhali:")
entry_denhali = tk.Entry(window)
label_debhali = tk.Label(window, text="Debhali:")
entry_debhali = tk.Entry(window)
label_ninhali = tk.Label(window, text="Ninhali:")
entry_ninhali = tk.Entry(window)


# Üret butonunu oluştur
generate_button = tk.Button(window, text="Üret", command=generate_docs)

# Sonuç etiketi
result_label = tk.Label(window, text="")

# Arayüz öğelerini düzenle
load_button.grid(row=0, column=0, columnspan=2)
label_ad.grid(row=1, column=0)
entry_ad.grid(row=1, column=1)
label_hikayeler.grid(row=2, column=0)
entry_hikayeler.grid(row=2, column=1)
label_cinsiyet.grid(row=3, column=0)
entry_cinsiyet.grid(row=3, column=1)
label_tip.grid(row=4, column=0)
entry_tip.grid(row=4, column=1)
generate_button.grid(row=5, column=0, columnspan=2)
result_label.grid(row=6, column=0, columnspan=2)
# İhali, Ehali, Dehali, Denhali, Debhali ve Ninhali etiketleri ve girdi alanlarını düzenle
label_ihali.grid(row=7, column=0)
entry_ihali.grid(row=7, column=1)
label_ehali.grid(row=8, column=0)
entry_ehali.grid(row=8, column=1)
label_dehali.grid(row=9, column=0)
entry_dehali.grid(row=9, column=1)
label_denhali.grid(row=10, column=0)
entry_denhali.grid(row=10, column=1)
label_debhali.grid(row=11, column=0)
entry_debhali.grid(row=11, column=1)
label_ninhali.grid(row=12, column=0)
entry_ninhali.grid(row=12, column=1)

window.mainloop()
